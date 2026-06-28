"""
Exportación de documentos a .docx para descarga.

El archivo original que sube el administrador no se almacena (solo se indexa
en pgvector). Para permitir la descarga, reconstruimos el texto desde los
chunks (ver atlas.get_document_text) y lo empaquetamos en un .docx generado
al vuelo con python-docx.
"""
import io
import re
from docx import Document

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def text_to_docx_bytes(title: str, text: str) -> bytes:
    """Genera un .docx en memoria a partir de un título y texto plano/markdown simple."""
    doc = Document()
    clean_title = re.sub(r"^\[Captura\]\s*", "", title).strip()
    doc.add_heading(clean_title or "Documento", level=0)

    for raw in (text or "").split("\n"):
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_filename(original_filename: str) -> str:
    """Convierte el nombre original (pdf/docx) a un nombre de descarga .docx."""
    base = re.sub(r"\.(pdf|docx|doc)$", "", original_filename or "documento", flags=re.IGNORECASE)
    base = re.sub(r"^\[Captura\]\s*", "", base).strip() or "documento"
    return f"{base}.docx"
